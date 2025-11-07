from fastapi import APIRouter, FastAPI, Request, HTTPException
import google.generativeai as genai
import os
import logging
from docx import Document
from datetime import datetime
import os
from methods.functions import Depends,require_role,get_db,Session
from models.schema import User,UserRole,RFP,Employee,Company
import marko

# Optional: for converting Word to PDF
from docx2pdf import convert  # You can replace this with another converter if needed

# Initialize FastAPI
app = FastAPI()

# Define API router
router = APIRouter(prefix="/api", tags=["RFP"])



@router.post("/going_to_edit", response_model=dict) #changes
def going_to_edit(rfp_data: dict,
    # current_user: User = Depends(require_role([UserRole.ADMIN, UserRole.EMPLOYEE])),
    db: Session = Depends(get_db)
):
    print("ulla vante")
    """Generate the final proposal document for the RFP"""
    print("adukum mela")
    print(rfp_data)
    company_id = rfp_data.get("company_id")
    rfp_id = rfp_data.get("rfp_id")
    employee_id = rfp_data.get("employee_id")
    print("id apro")
    print(employee_id)
    model_name = os.getenv("GEMINI_MODEL")
    model = None
    if model_name:
        try:
            model = genai.GenerativeModel(model_name)
        except Exception:
            logging.exception("Failed to initialize Gemini model '%s'", model_name)

    full_prompt = f"""
        You are an expert proposal writer. Compile the following question responses into a cohesive, professional
        proposal document that addresses the original RFP requirements.

        Format the proposal with appropriate sections, an executive summary, introduction, and conclusion.
        Ensure the document flows well and presents a compelling case for why our company should be selected.

        The final proposal should be in Markdown format with appropriate headings, bullet points, and formatting.

        rfp_data: {rfp_data}
        """

    if model is None:
        # Graceful fallback when there's no model configured
        logging.warning("No GEMINI_MODEL configured; returning unprocessed rfp_data as fallback result.")
        return {"prompt": "LLM not configured. Set GEMINI_MODEL to a supported model to enable proposal generation.", "rfp_data": rfp_data}

    prompt = model.generate_content(full_prompt)
    print("inga iruke")
    final_proposal_markdown = getattr(prompt, "text", str(prompt))
    print(final_proposal_markdown)
    print(company_id)
    # company = db.query(Company).filter(Company.id==company_id).first()
    # print(company_id)
    # print(company.subdomain)
    # if not company:
    #     return {"error":"company_id not found"}
    # subdomain = company.subdomain
    # # Generate Word and PDF documents from proposal
    # print("hi")
    # file_paths = generate_and_upload_proposal(company_id, {
    #     "title": "RFP Response",
    #     "final_proposal": final_proposal_markdown
    # },subdomain)
    # docx_url = file_paths.get("docx_url")
    # pdf_url = file_paths.get("pdf_url")
    # print(docx_url)
    # print(pdf_url)
    # rfp = db.query(RFP).filter(RFP.id==rfp_id).first()
    # rfp.status = "review pending"
    # if not rfp.docx_url:
    #     rfp.docx_url = docx_url
    # else:
    #     temp = rfp.docx_url
    #     rfp.docx_url = docx_url
    #     delete_s3_file(temp)
    # if not rfp.pdf_url:
    #     rfp.pdf_url = pdf_url
    # else:
    #     temp = rfp.docx_url
    #     rfp.pdf_url = pdf_url
    #     delete_s3_file(temp)
    # db.commit()
    # return file_paths
    return {"prompt": prompt.text}


@router.post("/final-rfp/status", response_model=dict)
async def final_rfp_status(request: Request, db: Session = Depends(get_db)):
    data = await request.json()
    rfp_id = data.get("rfp_id")
    if not rfp_id:
        return {"error": "rfp_id is required"}
    rfp = db.query(RFP).filter(RFP.id == rfp_id).first()
    if not rfp:
        return {"error": "RFP not found"}
    return {
        "docx_url": rfp.docx_url,
        "pdf_url": rfp.pdf_url
    }

# import boto3
# import os
# from io import BytesIO
# from datetime import datetime
# from docx import Document
# from docx2pdf import convert  # Only works on Windows with MS Word installed
# import tempfile
# import uuid

# def generate_and_upload_proposal(company_id, responses,subdomain):
#     # Setup document
#     doc = Document()
#     title = responses.get("title", "RFP Response")
#     doc.add_heading(title, 0)
#     doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}\n")

#     if final_proposal := responses.get("final_proposal", ""):
#         sections = final_proposal.split("## ")
#         for i, section in enumerate(sections):
#             if i == 0 and not section.startswith("#"):
#                 doc.add_paragraph(section)
#                 continue
#             lines = section.split("\n", 1)
#             heading = lines[0].strip()
#             doc.add_heading(heading, level=2)
#             if len(lines) > 1:
#                 for para in lines[1].split("\n\n"):
#                     para = para.strip()
#                     if para.startswith("- "):
#                         for item in para.split("- "):
#                             if item.strip():
#                                 doc.add_paragraph(item.strip(), style='List Bullet')
#                     elif para:
#                         doc.add_paragraph(para)
#     else:
#         for q_id, response in responses.items():
#             if q_id not in ["final_proposal", "title"]:
#                 doc.add_heading(f"Question {q_id}", level=2)
#                 doc.add_paragraph(response)

#     # Save .docx to temp
#     timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
#     filename_base = f"{title.replace(' ', '_')}_{timestamp}"
#     with tempfile.TemporaryDirectory() as tmpdir:
#         docx_path = os.path.join(tmpdir, f"{filename_base}.docx")
#         pdf_path = os.path.join(tmpdir, f"{filename_base}.pdf")
#         doc.save(docx_path)

#         # Convert to PDF
#         convert(docx_path, pdf_path)

#         AWS_ACCESS_KEY_ID = os.getenv("ACCESS_KEY_AWS")
#         AWS_SECRET_ACCESS_KEY = os.getenv("SECRET_KEY_AWS")
#         BUCKET_NAME = "rfp-storage-bucket"
#         REGION = "us-east-1"  # use your S3 bucket region
#         # Upload both files to S3
#         s3 = boto3.client(
#             "s3",
#             region_name=REGION,
#             aws_access_key_id=AWS_ACCESS_KEY_ID,
#             aws_secret_access_key=AWS_SECRET_ACCESS_KEY
#         )
        
#         docx_key = f"proposals/{uuid.uuid4()}.{subdomain}.docx"
#         pdf_key = f"proposals/{uuid.uuid4()}.{subdomain}.pdf"

#         with open(docx_path, "rb") as docx_file:
#             s3.upload_fileobj(docx_file, BUCKET_NAME, docx_key, ExtraArgs={"ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document","ContentDisposition": "inline"}
            
#         ) #multipart-data

#         with open(pdf_path, "rb") as pdf_file:
#             s3.upload_fileobj(pdf_file, BUCKET_NAME, pdf_key, ExtraArgs={"ContentType": "application/pdf"})

#         docx_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{docx_key}"
#         pdf_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{pdf_key}"

#         with open(pdf_path, "rb") as pdf_file:
#             s3.upload_fileobj(pdf_file, BUCKET_NAME, pdf_key, ExtraArgs={"ContentType": "application/pdf"})

#         docx_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{docx_key}"
#         pdf_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{pdf_key}"

#     return {
#         "status": "success",
#         "docx_url": docx_url,
#         "pdf_url": pdf_url
#     }

#     return {
#         "status": "success",
#         "docx_url": docx_url,
#         "pdf_url": pdf_url
#     }