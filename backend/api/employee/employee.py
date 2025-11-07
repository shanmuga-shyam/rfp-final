# employee.py
from models.schema import User, UserRole, UserCreate, UserResponse, RFP, Employee, EmployeeCreate, Company
from methods.functions import get_db, require_role, get_password_hash
from methods.functions import Session
from fastapi import HTTPException, Depends, Form
from fastapi import APIRouter
from typing import List, Optional
from fastapi.responses import StreamingResponse
from fastapi import status, Request
from datetime import datetime
import os
import logging
import json
from pydantic import BaseModel
import google.generativeai as genai

router = APIRouter(prefix="/api", tags=["Employee"])

# In-memory dictionary to store current_rfp_id for each employee
# {employee_id: rfp_id}
current_rfp_ids_in_memory: dict[int, int] = {}
current_rfp_context_in_memory: dict[int, dict] = {} # {employee_id: {rfp_id, company_id, filename, file_url}}


class AssignedRFPResponse(BaseModel):
    id: int
    filename: str
    created_at: str
    updated_at: Optional[str] = None
    status: str
    content_type: str
    file_url: str # Ensure file_url is part of the response model

@router.get("/employee/get_assigned_rfps/{employee_id}", response_model=dict)
async def get_assigned_rfps(employee_id: int, db: Session = Depends(get_db)):
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        return {"rfps": []}
    
    assigned_ids = employee.rfps_assigned
    
    # Handle both JSON string and already parsed list
    if isinstance(assigned_ids, str):
        try:
            assigned_ids = json.loads(assigned_ids)
        except json.JSONDecodeError:
            assigned_ids = []
    elif assigned_ids is None:
        assigned_ids = []

    arr = []
    # Fetch RFP details for each assigned ID
    rfp_files = db.query(RFP).filter(RFP.id.in_(assigned_ids)).all()
    for rfpfile in rfp_files:
        arr.append({
            "id": rfpfile.id,
            "filename": rfpfile.filename,
            "content_type": rfpfile.content_type,
            "status": rfpfile.status,
            "created_at": rfpfile.created_at.isoformat() if rfpfile.created_at else None,
            "file_url": rfpfile.file_url # Include file_url
        })
    return {"rfps": arr}


@router.post('/set_current_rfp/{employee_id}')
def set_current_rfp(
    employee_id: int,
    rfp_id: int,
    db: Session = Depends(get_db)
):
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="Employee not found")

    # Store in the in-memory variable
    current_rfp_ids_in_memory[employee_id] = rfp_id
    print(f"Stored in memory: Employee {employee_id} -> RFP {rfp_id}")
    return {"message": "Current RFP ID set successfully (in-memory)", "current_rfp_id": rfp_id}


@router.get("/employee/completed_rfps/{employee_id}")
def get_completed_rfps_by_employee(employee_id: int, db: Session = Depends(get_db)):
    """
    Returns all RFPs completed by the employee (using rfps_finished column).
    """
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        print(f"No employee found for id {employee_id}")
        return {"rfps": []}
    finished_ids = getattr(employee, 'rfps_finished', None)
    print(f"Raw rfps_finished for employee {employee_id}: {finished_ids}")
    if isinstance(finished_ids, str):
        try:
            finished_ids = json.loads(finished_ids)
        except Exception as e:
            print(f"Error parsing rfps_finished: {e}")
            finished_ids = []
    if not isinstance(finished_ids, list):
        finished_ids = []
    print(f"Parsed rfps_finished for employee {employee_id}: {finished_ids}")
    arr = []
    if finished_ids:
        rfp_files = db.query(RFP).filter(RFP.id.in_(finished_ids)).all()
        print(f"Found {len(rfp_files)} finished RFPs for employee {employee_id}")
        for rfpfile in rfp_files:
            arr.append({
                "filename": getattr(rfpfile, "filename", None)
            })
    print(f"Response for finished RFPs: {arr}")
    return {"rfps": arr}

@router.get("/employee/company_name/{employee_id}")
def get_company_name_by_employee(employee_id: int, db: Session = Depends(get_db)):
    """
    Returns the company name for a given employee id.
    """
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        return {"company_name": None}
    company_id = getattr(employee, 'company_id', None)
    if not company_id:
        return {"company_name": None}
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        return {"company_name": None}
    return {"company_name": company.name}

import requests
import os
import tempfile
from agents.extract_rfp_structure import process_document

@router.post("/employee/final_rfp") #changes
def final_rfp(
    text: str = Form(...), 
    changes: str = Form(...),
    db: Session = Depends(get_db)
):
    # Use GEMINI_MODEL env var for model name; instantiate lazily and handle failure
    model_name = os.getenv("GEMINI_MODEL")
    model = None
    if model_name:
        try:
            model = genai.GenerativeModel(model_name)
        except Exception:
            logging.exception("Failed to initialize Gemini model '%s'", model_name)

    if model is None:
        # Graceful error telling the operator to configure the model
        raise HTTPException(
            status_code=503,
            detail=(
                "LLM not configured or unavailable. Set GEMINI_MODEL to a supported model name and ensure API credentials."
            ),
        )

    try:
        prompt = model.generate_content(
            f"""
            I have a document which contains the text "{text}". I want you to apply the following {changes} to each relevant part of the data. Modify the content accordingly and return the final output in the correct order, preserving structure and formatting. Apply only the changes mentioned—do not invent or omit anything.
            """
        )
        print(getattr(prompt, "text", str(prompt)))
        return {"prompt": getattr(prompt, "text", str(prompt))}
    except Exception:
        logging.exception("LLM generate_content failed in final_rfp")
        raise HTTPException(status_code=500, detail="LLM generation failed; check server logs.")


import boto3
from botocore.exceptions import ClientError
from urllib.parse import urlparse

from dotenv import load_dotenv
load_dotenv()
def parse_s3_url(s3_url: str):
    """Extract bucket name and key from S3 URL."""
    parsed = urlparse(s3_url)
    bucket = parsed.netloc
    key = parsed.path.lstrip('/')
    return bucket, key

def delete_s3_file(s3_url: str):
    try:
        bucket_name, key = parse_s3_url(s3_url)
        s3_client = boto3.client('s3')
        s3_client.delete_object(Bucket=bucket_name, Key=key)
        return True
    except Exception as e:
        print(f"Error deleting file: {e}")
        return False
    
@router.post("/employee/ok")
def ok(
    text: str = Form(...),
    rfp_id:str = Form(...),
    db: Session = Depends(get_db)
):
    print("ok kula iruke")
    rfp = db.query(RFP).filter(RFP.id==int(rfp_id)).first()
    print(rfp)
    if not rfp: # Added check for RFP
        raise HTTPException(status_code=404, detail="RFP not found")
    
    
    company = db.query(Company).filter(Company.id==rfp.company_id).first()
    print(company.subdomain)
    if not company:
        return {"error":"company_id not found"}
    subdomain = company.subdomain
    # Generate Word and PDF documents from proposal
    print("hi")
    file_paths = generate_and_upload_proposal(company.id, {
        "title": "RFP Response",
        "final_proposal": text
    },subdomain)
    
    docx_url = file_paths.get("docx_url")
    pdf_url = file_paths.get("pdf_url")
    print(docx_url)
    print(pdf_url)
    
    rfp.status = "review pending"
    # DOCX replacement and deletion
    if not rfp.docx_url:
        rfp.docx_url = docx_url
    else:
        old_docx = rfp.docx_url
        rfp.docx_url = docx_url
        delete_s3_file(old_docx)

    # PDF replacement and deletion
    if not rfp.pdf_url:
        rfp.pdf_url = pdf_url
    else:
        old_pdf = rfp.pdf_url
        rfp.pdf_url = pdf_url
        delete_s3_file(old_pdf)

    db.commit()
    return file_paths
    

def generate_and_upload_proposal(company_id, responses, subdomain):
    from docx import Document
    from docx.shared import Pt
    from docx2pdf import convert
    import tempfile
    import os
    import uuid
    import boto3
    from datetime import datetime
    import re

    # Setup document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title and Date
    title = responses.get("title", "RFP Response")
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")

    final_proposal = responses.get("final_proposal", "")

    if final_proposal:
        # Clean up and normalize line endings
        final_proposal = final_proposal.replace("\\n", "\n").strip()

        # Split sections by markdown header '##'
        sections = re.split(r"\n##\s*", final_proposal)

        for section in sections:
            if not section.strip():
                continue

            lines = section.strip().split("\n", 1)
            heading = lines[0].strip()
            doc.add_heading(heading, level=2)

            if len(lines) > 1:
                body = lines[1].strip()
                paragraphs = body.split("\n\n")
                for para in paragraphs:
                    para = para.strip()
                    if not para:
                        continue

                    # Bullet points
                    if para.startswith("* "):
                        bullets = para.split("\n")
                        for bullet in bullets:
                            if bullet.startswith("* "):
                                run = doc.add_paragraph(style="List Bullet").add_run(bullet[2:].strip())
                                run.font.name = "Calibri"
                                run.font.size = Pt(11)
                    else:
                        # Bold formatting
                        parts = re.split(r"(\*\*.*?\*\*)", para)
                        p = doc.add_paragraph()
                        for part in parts:
                            if part.startswith("**") and part.endswith("**"):
                                run = p.add_run(part[2:-2])
                                run.bold = True
                            else:
                                p.add_run(part)
    else:
        # Fallback: Add individual questions if no final_proposal
        for q_id, response in responses.items():
            if q_id not in ["final_proposal", "title"]:
                doc.add_heading(f"Question {q_id}", level=2)
                doc.add_paragraph(response)

    # Remove empty first paragraph if any
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    # Save to temp paths
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename_base = f"{title.replace(' ', '_')}_{timestamp}"
    unique_id = str(uuid.uuid4())

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, f"{filename_base}.docx")
        pdf_path = os.path.join(tmpdir, f"{filename_base}.pdf")
        doc.save(docx_path)

        try:
            convert(docx_path, pdf_path)
        except Exception as e:
            return {
                "status": "error",
                "message": f"PDF conversion failed: {e}"
            }

        # AWS S3 Upload
        AWS_ACCESS_KEY_ID = os.getenv("ACCESS_KEY_AWS")
        AWS_SECRET_ACCESS_KEY = os.getenv("SECRET_KEY_AWS")
        BUCKET_NAME = "rfp-storage-bucket"
        REGION = "us-east-1"

        s3 = boto3.client(
            "s3",
            region_name=REGION,
            aws_access_key_id=AWS_ACCESS_KEY_ID,
            aws_secret_access_key=AWS_SECRET_ACCESS_KEY
        )

        docx_key = f"proposals/{unique_id}.{subdomain}.docx"
        pdf_key = f"proposals/{unique_id}.{subdomain}.pdf"

        try:
            with open(docx_path, "rb") as docx_file:
                s3.upload_fileobj(
                    docx_file,
                    BUCKET_NAME,
                    docx_key,
                    ExtraArgs={
                        "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "ContentDisposition": "attachment"
                    }
                )
            with open(pdf_path, "rb") as pdf_file:
                s3.upload_fileobj(
                    pdf_file,
                    BUCKET_NAME,
                    pdf_key,
                    ExtraArgs={"ContentType": "application/pdf"}
                )
        except Exception as e:
            return {
                "status": "error",
                "message": f"Failed to upload to S3: {e}"
            }

        docx_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{docx_key}"
        pdf_url = f"https://{BUCKET_NAME}.s3.amazonaws.com/{pdf_key}"

    return {
        "status": "success",
        "docx_url": docx_url,
        "pdf_url": pdf_url
    }

@router.get("/employee/rfps/{rfp_id}/response")
def get_rfp_response(rfp_id: int, db: Session = Depends(get_db)):
    rfp = db.query(RFP).filter(RFP.id == rfp_id).first()
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
    # Try to get the generated response (adjust the attribute as per your model)
    response_data = getattr(rfp, "generated_response", None) or getattr(rfp, "response_json", None)
    if response_data:
        # If the response is a stringified JSON, parse it
        try:
            import json
            response_json = json.loads(response_data) if isinstance(response_data, str) else response_data
            return {"response": response_json}
        except Exception:
            return {"response": response_data}
    # Fallback: return extracted text if available
    extracted_text = getattr(rfp, "extracted_text", None)
    if extracted_text:
        return {"extracted_text": extracted_text}
    return {"message": "No response or extracted text found for this RFP."}
