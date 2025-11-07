import requests
from fastapi import APIRouter, Depends, HTTPException, Request
from sqlalchemy.orm import Session
from models.schema import RFP, Employee, Company
from methods.functions import get_db
from langchain_groq import ChatGroq
from langchain.agents import initialize_agent, AgentType
from agents.tools.company_doc_tool import get_company_qa_tool
from agents.tools.fall_back_tool import FallbackLLMTool
import os
from PyPDF2 import PdfReader
from io import BytesIO
import google.generativeai as genai
import os
import logging

router = APIRouter(prefix="/api", tags=["EmployeeProposalEdit"])

@router.post("/employee/rfps/{rfp_id}/extract-pdf-text")
def extract_pdf_text(
    rfp_id: int,
    db: Session = Depends(get_db)
):
    rfp = db.query(RFP).filter(RFP.id == rfp_id).first()
    if not rfp or not rfp.pdf_url:
        raise HTTPException(status_code=404, detail="RFP or PDF not found.")
    try:
        response = requests.get(rfp.pdf_url)
        response.raise_for_status()
        pdf_bytes = BytesIO(response.content)
        reader = PdfReader(pdf_bytes)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract PDF text: {str(e)}")
    return {"text": text}

@router.post("/employee/rfps/{rfp_id}/extract-file-text")
def extract_file_text(
    rfp_id: int,
    db: Session = Depends(get_db)
):
    print("hi")
    print(f"[extract-file-text] Endpoint called for rfp_id={rfp_id}")
    rfp = db.query(RFP).filter(RFP.id == rfp_id).first()
    if not rfp or (not rfp.pdf_url and not rfp.docx_url):
        print(f"[extract-file-text] RFP not found or missing file url. RFP: {rfp}")
        raise HTTPException(status_code=404, detail="RFP or file not found.")
    try:
        print(f"[extract-file-text] RFP: {rfp}")
        if rfp.pdf_url:
            print(f"[extract-file-text] Trying PDF url: {rfp.pdf_url}")
            response = requests.get(rfp.pdf_url)
            response.raise_for_status()
            pdf_bytes = BytesIO(response.content)
            reader = PdfReader(pdf_bytes)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif rfp.docx_url:
            import docx
            print(f"[extract-file-text] Trying DOCX url: {rfp.docx_url}")
            response = requests.get(rfp.docx_url)
            response.raise_for_status()
            docx_bytes = BytesIO(response.content)
            doc = docx.Document(docx_bytes)
            text = "\n".join([p.text for p in doc.paragraphs])
        else:
            print("[extract-file-text] No file url present on RFP.")
            text = ""
    except Exception as e:
        print(f"[extract-file-text] Exception: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to extract file text: {str(e)}")
    return {"text": text}

@router.post("/employee/rfps/{rfp_id}/custom-prompt-edit")
async def employee_custom_prompt_edit(
    rfp_id: int,
    payload: dict,
    db: Session = Depends(get_db)
):
    prompt = payload.get("prompt")
    if not prompt:
        print(f"[custom-prompt-edit] No prompt provided.")
        raise HTTPException(status_code=400, detail="Prompt is required.")
    rfp = db.query(RFP).filter(RFP.id == rfp_id).first()
    if not rfp or (not rfp.pdf_url and not rfp.docx_url):
        print(f"[custom-prompt-edit] RFP not found or missing file url. RFP: {rfp}")
        raise HTTPException(status_code=404, detail="RFP or file not found.")
    # Extract file text (reuse logic)
    try:
        print(f"[custom-prompt-edit] RFP: {rfp}")
        if rfp.pdf_url:
            print(f"[custom-prompt-edit] Trying PDF url: {rfp.pdf_url}")
            response = requests.get(rfp.pdf_url)
            response.raise_for_status()
            pdf_bytes = BytesIO(response.content)
            reader = PdfReader(pdf_bytes)
            file_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif rfp.docx_url:
            import docx
            print(f"[custom-prompt-edit] Trying DOCX url: {rfp.docx_url}")
            response = requests.get(rfp.docx_url)
            response.raise_for_status()
            docx_bytes = BytesIO(response.content)
            doc = docx.Document(docx_bytes)
            file_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            print("[custom-prompt-edit] No file url present on RFP.")
            file_text = ""
    except Exception as e:
        print(f"[custom-prompt-edit] Exception during extraction: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to extract file text: {str(e)}")
    # LLM prompt
    try:
        model_name = os.getenv("GEMINI_MODEL")
        model = None
        if model_name:
            try:
                model = genai.GenerativeModel(model_name)
            except Exception:
                logging.exception("Failed to initialize Gemini model '%s'", model_name)

        full_prompt = f"File Content:\n{file_text}\n\nInstruction: {prompt}"
        if model is None:
            # graceful fallback when model is not configured
            result = (
                "LLM not configured or unavailable. Set the GEMINI_MODEL environment variable to a supported model "
                "and ensure the API key/permissions are correct."
            )
        else:
            try:
                response = model.generate_content(full_prompt)
                result = getattr(response, 'text', str(response))
            except Exception:
                logging.exception("LLM generate_content failed")
                result = "LLM call failed; check server logs for details."
    except Exception as e:
        logging.exception("Unexpected error in custom prompt edit")
        result = f"Error: {str(e)}"
    return {"result": result}

@router.post("/employee/rfps/{rfp_id}/final-proposal")
async def employee_final_proposal(rfp_id: int, request: Request, db: Session = Depends(get_db)):
    """
    Generate the final proposal for an employee's RFP using Gemini LLM and return the result.
    Expects JSON body: { "proposal": <string> }
    """
    try:
        data = await request.json()
        proposal_text = data.get("proposal")
        if not proposal_text:
            raise HTTPException(status_code=400, detail="Proposal text is required.")
        rfp = db.query(RFP).filter(RFP.id == rfp_id).first()
        if not rfp:
            raise HTTPException(status_code=404, detail="RFP not found.")
        # Optionally, you can add more context from the RFP or employee here
        model_name = os.getenv("GEMINI_MODEL")
        model = None
        if model_name:
            try:
                model = genai.GenerativeModel(model_name)
            except Exception:
                logging.exception("Failed to initialize Gemini model '%s'", model_name)

        prompt_text = (
            f"You are an expert proposal writer. Refine and finalize the following proposal draft into a professional, cohesive document suitable for submission. Format with appropriate sections, summary, and conclusion. Return the result in Markdown format.\n\nProposal Draft:\n{proposal_text}"
        )

        if model is None:
            # graceful fallback
            return {
                "result": (
                    "LLM not configured or unavailable. Set the GEMINI_MODEL environment variable to a supported model "
                    "and ensure API credentials are available."
                )
            }

        try:
            prompt = model.generate_content(prompt_text)
            final_proposal_markdown = getattr(prompt, 'text', str(prompt))
            return {"result": final_proposal_markdown}
        except Exception:
            logging.exception("LLM generate_content failed for final proposal")
            raise HTTPException(status_code=500, detail="LLM generation failed; check server logs.")
    except Exception as e:
        logging.exception("Unexpected error in employee_final_proposal")
        raise HTTPException(status_code=500, detail=f"Failed to generate final proposal: {str(e)}")
